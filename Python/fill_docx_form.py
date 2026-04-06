import re
from docx import Document
from docx.shared import RGBColor

def extract_docx_fields(docx_path):
    doc = Document(docx_path)
    fields = []
    seen = set()

    for t_index, table in enumerate(doc.tables):
        for r_index, row in enumerate(table.rows):
            for c_index, cell in enumerate(row.cells):
                text = cell.text.strip()
                if not text:
                    continue
                parts = re.split(r'\n+|\t+| {2,}', text)
                for part in parts:
                    field = part.strip()
                    if not field:
                        continue
                    field = field.replace("*", "").strip()
                    field = field.rstrip(":")
                    if len(field) > 60:
                        continue
                    if field.lower() in [
                        "supplier information form",
                        "general information",
                        "business specific information",
                        "part i","part ii","part iii","part iv","part v","part vi"
                    ]:
                        continue

                    if field in seen:
                        continue
                    seen.add(field)
                    fields.append({
                        "field_name": field,
                        "location": ("table", t_index, r_index, c_index),
                        "field_type": "text"
                    })

    return fields, doc

def autofill_docx(doc, fields, match_results, output_path):

    match_lookup = {m.get("form_question"): m for m in match_results}

    for f in fields:
        field_name = f["field_name"]
        location = f["location"]

        match = match_lookup.get(field_name)

        if not match:
            continue

        decision = match.get("decision", "").lower()

        if decision not in ["autofill", "review suggested"]:
            continue

        answer = str(match.get("answer", ""))

        table_type, t_index, r_index, c_index = location

        table = doc.tables[t_index]

        target_col = min(c_index + 1, len(table.rows[r_index].cells) - 1)
        cell = table.rows[r_index].cells[target_col]

        cell.text = answer

        if decision == "review suggested":
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.highlight_color = 7

    doc.save(output_path)

    print(f"DOCX saved as {output_path}")